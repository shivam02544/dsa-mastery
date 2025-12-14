import os
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run 'pip install python-docx'")
    exit(1)

def create_document():
    document = Document()

    # Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = document.add_heading('DSA Mastery Project Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Student Info
    p = document.add_paragraph('Submitted by: [Your Info]\nSubject: Data Structures and Algorithms')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Spacer

    # 1. Aim
    document.add_heading('1. Aim', level=1)
    document.add_paragraph(
        "To create an interactive and student-friendly platform, 'DSA Mastery', that helps users visualize and "
        "understand complex Data Structures and Algorithms (DSA) concepts through real-time animations, "
        "while tracking their learning progress."
    )

    # 2. Description
    document.add_heading('2. Description', level=1)
    document.add_paragraph(
        "DSA Mastery is a Web Application designed to bridge the gap between theoretical DSA concepts and "
        "practical implementation. It features a robust Sorting Visualizer that demonstrates algorithms like "
        "Bubble Sort, Quick Sort, and Merge Sort step-by-step. The platform includes a user profile system for "
        "tracking progress, powered by MongoDB, and potentially Google Authentication. "
        "The user interface is built with Next.js and Tailwind CSS, focusing on a modern, responsive, "
        "and engaging design with glassmorphism effects."
    )

    # 3. Design
    document.add_heading('3. What is the Design?', level=1)
    design_p = document.add_paragraph()
    design_p.add_run("Tech Stack: ").bold = True
    design_p.add_run("Next.js 14, Tailwind CSS, MongoDB, NextAuth.\n")
    design_p.add_run("User Interface: ").bold = True
    design_p.add_run("Modern 'Glassmorphism' aesthetic with dark mode, smooth framer-motion animations, and responsive layouts.\n")
    design_p.add_run("Core Logic: ").bold = True
    design_p.add_run("React Hooks (useState, useEffect) manage the visualizer state, where array values are mapped to bar heights and colors indicate state changes (Compare/Swap).")

    # 4. Implementation Code
    document.add_heading('4. Implementation Code', level=1)
    document.add_paragraph("Below are key snippets from the project source code:")

    document.add_heading('SortingVisualizer.jsx (Core Logic)', level=2)
    
    code_path = os.path.join('components', 'SortingVisualizer.jsx')
    if os.path.exists(code_path):
        with open(code_path, 'r', encoding='utf-8') as f:
            code_content = f.read()
            # Truncate for brevity in report
            lines = code_content.split('\n')
            snippet = '\n'.join(lines[0:150]) + "\n\n// ... (Methods for other algorithms like Quick/Merge Sort) ...\n"
            
            p = document.add_paragraph(snippet)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
    else:
        document.add_paragraph("(SortingVisualizer.jsx not found)")

    document.add_page_break()

    document.add_heading('App/Page.js (Home Design)', level=2)
    home_path = os.path.join('app', 'page.js')
    if os.path.exists(home_path):
        with open(home_path, 'r', encoding='utf-8') as f:
            code_content = f.read()
            p = document.add_paragraph(code_content)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
    else:
        document.add_paragraph("(app/page.js not found)")

    # 5. Results (Screenshots)
    document.add_heading('5. Results', level=1)
    document.add_paragraph("The following screenshots demonstrate the application in action:")

    def add_screenshot(filename, caption):
        path = os.path.join('public', 'report_screenshots', filename)
        if os.path.exists(path):
            document.add_heading(caption, level=3)
            try:
                document.add_picture(path, width=Inches(6.0))
            except Exception as e:
                document.add_paragraph(f"[Error adding image: {e}]")
            document.add_paragraph("\n")
        else:
            document.add_paragraph(f"[Screenshot missing: {caption} - File not found at {path}]")

    add_screenshot('home.png', 'Home Page')
    add_screenshot('visualizer.png', 'Sorting Visualizer Action')
    add_screenshot('profile.png', 'User Profile Page')

    output_file = 'DSA_Mastery_Project_Report.docx'
    document.save(output_file)
    print(f"Document created successfully: {output_file}")

if __name__ == "__main__":
    create_document()
